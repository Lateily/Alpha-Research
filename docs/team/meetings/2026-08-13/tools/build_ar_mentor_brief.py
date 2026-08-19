from __future__ import annotations

from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


OUT_DIR = Path(r"C:\Users\jctx\Desktop\AR\output\meetings")
OUT_PATH = OUT_DIR / "2026-08-15_AR_第三周至今_学长会议系统汇报.docx"
FIG_PATH = OUT_DIR / "2026-08-15_AR_三层系统与产品落地图.png"

BLUE = "2E74B5"
DEEP_BLUE = "1F4D78"
NAVY = "17365D"
LIGHT_BLUE = "D9EAF7"
PALE_BLUE = "EEF5FA"
LIGHT_GRAY = "F2F4F7"
MID_GRAY = "D9DEE7"
DARK = "1F2937"
MUTED = "5B6573"
GREEN = "2E7D32"
GREEN_BG = "E8F5E9"
AMBER = "9A6700"
AMBER_BG = "FFF4CE"
RED = "B42318"
RED_BG = "FDECEC"
WHITE = "FFFFFF"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color="D5DAE2", size="6") -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), size)
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), color)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_width(table, width_dxa=9360) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_dxa))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")


def set_table_grid(table, widths: list[int]) -> None:
    tbl_grid = table._tbl.tblGrid
    for child in list(tbl_grid):
        tbl_grid.remove(child)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        tbl_grid.append(grid_col)


def set_font(run, size=None, bold=None, color=None, name="Calibri", east_asia="Microsoft YaHei"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_text(paragraph, text: str, *, bold=False, color=None, size=None, italic=False, code=False):
    run = paragraph.add_run(text)
    set_font(run, size=size, bold=bold, color=color,
             name="Consolas" if code else "Calibri",
             east_asia="Microsoft YaHei")
    run.italic = italic
    return run


def add_hyperlink(paragraph, text: str, url: str, color=BLUE):
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:ascii"), "Calibri")
    r_fonts.set(qn("w:hAnsi"), "Calibri")
    r_fonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    r_pr.append(r_fonts)
    c = OxmlElement("w:color")
    c.set(qn("w:val"), color)
    r_pr.append(c)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(underline)
    new_run.append(r_pr)
    text_node = OxmlElement("w:t")
    text_node.text = text
    new_run.append(text_node)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def add_bottom_border(paragraph, color=BLUE, size="18", space="4") -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), space)
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_text(paragraph, "AR · 学长会议汇报  |  ", size=8.5, color=MUTED)
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)
    set_font(run, size=8.5, color=MUTED)


def configure_styles(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(DARK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in (
        ("Title", 24, NAVY, 0, 8),
        ("Subtitle", 12, MUTED, 0, 12),
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 11.5, DEEP_BLUE, 8, 4),
    ):
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = name != "Subtitle"
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for sname in ("List Bullet", "List Number"):
        st = styles[sname]
        st.font.name = "Calibri"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        st.font.size = Pt(10.5)
        st.paragraph_format.left_indent = Inches(0.5)
        st.paragraph_format.first_line_indent = Inches(-0.25)
        st.paragraph_format.space_after = Pt(5)
        st.paragraph_format.line_spacing = 1.10

    if "Small Note" not in styles:
        note = styles.add_style("Small Note", WD_STYLE_TYPE.PARAGRAPH)
    else:
        note = styles["Small Note"]
    note.font.name = "Calibri"
    note._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    note.font.size = Pt(8.5)
    note.font.color.rgb = RGBColor.from_string(MUTED)
    note.paragraph_format.space_after = Pt(4)
    note.paragraph_format.line_spacing = 1.0


def set_section(section) -> None:
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.78)
    section.bottom_margin = Inches(0.70)
    section.left_margin = Inches(0.84)
    section.right_margin = Inches(0.84)
    section.header_distance = Inches(0.32)
    section.footer_distance = Inches(0.35)

    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_text(hp, "INTERNAL · CURRENT-STATE BRIEF · 2026-08-15", size=8, color=MUTED)

    footer = section.footer
    add_page_number(footer.paragraphs[0])


def add_heading(doc: Document, text: str, level=1):
    return doc.add_heading(text, level=level)


def add_bullet(doc: Document, text: str, *, bold_prefix: str | None = None, level=0):
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    if bold_prefix and text.startswith(bold_prefix):
        add_text(p, bold_prefix, bold=True, color=DEEP_BLUE)
        add_text(p, text[len(bold_prefix):])
    else:
        add_text(p, text)
    return p


def add_number(doc: Document, text: str):
    p = doc.add_paragraph(style="List Number")
    add_text(p, text)
    return p


def add_callout(doc: Document, title: str, body: str, *, fill=PALE_BLUE, accent=BLUE, status: str | None = None):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table, 9360)
    set_table_grid(table, [9360])
    set_table_borders(table, color=accent, size="10")
    cell = table.cell(0, 0)
    set_cell_width(cell, 9360)
    set_cell_shading(cell, fill)
    set_cell_margins(cell, 150, 120, 150, 120)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    add_text(p, title, bold=True, color=accent, size=11)
    if status:
        add_text(p, f"  [{status}]", bold=True, color=accent, size=9)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    add_text(p2, body, size=10)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_status_table(doc: Document, headers: list[str], rows: Iterable[Iterable[str]], widths: list[int], status_col: int | None = None, font_size=9.2):
    rows = [list(r) for r in rows]
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table, sum(widths))
    set_table_grid(table, widths)
    set_table_borders(table)
    table.autofit = False
    header = table.rows[0]
    set_repeat_table_header(header)
    for i, value in enumerate(headers):
        cell = header.cells[i]
        set_cell_width(cell, widths[i])
        set_cell_shading(cell, LIGHT_GRAY)
        set_cell_margins(cell)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_after = Pt(0)
        add_text(p, value, bold=True, color=NAVY, size=font_size)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for ridx, row_data in enumerate(rows):
        row = table.add_row()
        prevent_row_split(row)
        if ridx % 2 == 1:
            for c in row.cells:
                set_cell_shading(c, "FAFBFC")
        for i, value in enumerate(row_data):
            cell = row.cells[i]
            set_cell_width(cell, widths[i])
            set_cell_margins(cell)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            color = DARK
            bold = False
            if status_col is not None and i == status_col:
                bold = True
                v = value.upper()
                if any(k in v for k in ("DONE", "MERGED", "已合并", "主干", "GREEN")):
                    set_cell_shading(cell, GREEN_BG)
                    color = GREEN
                elif any(k in v for k in ("BLOCK", "RED", "未启动", "未接线")):
                    set_cell_shading(cell, RED_BG)
                    color = RED
                else:
                    set_cell_shading(cell, AMBER_BG)
                    color = AMBER
            add_text(p, str(value), bold=bold, color=color, size=font_size)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def page_break(doc: Document):
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def find_font(size: int, bold=False):
    candidates = [
        r"C:\Windows\Fonts\msyhbd.ttc" if bold else r"C:\Windows\Fonts\msyh.ttc",
        r"C:\Windows\Fonts\simhei.ttf" if bold else r"C:\Windows\Fonts\simsun.ttc",
        r"C:\Windows\Fonts\arialbd.ttf" if bold else r"C:\Windows\Fonts\arial.ttf",
    ]
    for path in candidates:
        if Path(path).exists():
            try:
                return ImageFont.truetype(path, size=size)
            except OSError:
                pass
    return ImageFont.load_default()


def draw_centered(draw, box, text, font, fill):
    x0, y0, x1, y1 = box
    bbox = draw.multiline_textbbox((0, 0), text, font=font, spacing=6, align="center")
    w = bbox[2] - bbox[0]
    h = bbox[3] - bbox[1]
    draw.multiline_text(((x0 + x1 - w) / 2, (y0 + y1 - h) / 2), text, font=font, fill=fill, spacing=6, align="center")


def make_architecture_figure(path: Path) -> None:
    W, H = 1800, 1000
    img = Image.new("RGB", (W, H), "#FFFFFF")
    d = ImageDraw.Draw(img)
    title_font = find_font(44, True)
    band_font = find_font(31, True)
    small_font = find_font(23, False)
    tiny_font = find_font(20, False)
    d.text((70, 45), "AR 从研究能力到产品价值的主链", font=title_font, fill="#17365D")
    d.text((70, 105), "事实与证据先行；AIOS 负责调度与自证；产品层负责让人完成工作", font=small_font, fill="#5B6573")

    layers = [
        ("产品体验层", "驾驶舱｜市场研究｜宏观｜组合｜复盘｜周报｜工作台", "#D9EAF7", "#17365D"),
        ("AIOS Harness", "任务契约 → 权威上下文 → Policy → Router/Agent → Evals/证据 → Memory/观测", "#FFF4CE", "#7A5000"),
        ("Research OS", "U0 全市场注册 → U1 扫描 → U2 候选/随机对照 → U3 六维电池 → U4 人审 → U5 决策/验证", "#E8F5E9", "#205B25"),
        ("工程与数据底座", "Nightly｜Point-in-time 数据｜事件账本｜Macro OS｜API/Contracts｜GitHub/CI/Worktrees", "#F2F4F7", "#374151"),
    ]
    y = 185
    layer_h = 142
    for idx, (name, desc, bg, fg) in enumerate(layers):
        d.rounded_rectangle((70, y, 1730, y + layer_h), radius=22, fill=bg, outline=fg, width=3)
        d.rounded_rectangle((94, y + 26, 390, y + layer_h - 26), radius=18, fill=fg)
        draw_centered(d, (94, y + 26, 390, y + layer_h - 26), name, band_font, "#FFFFFF")
        draw_centered(d, (420, y + 18, 1695, y + layer_h - 18), desc, small_font, fg)
        if idx < len(layers) - 1:
            cx = 900
            d.line((cx, y + layer_h + 4, cx, y + layer_h + 25), fill="#8391A2", width=5)
            d.polygon([(cx - 10, y + layer_h + 18), (cx + 10, y + layer_h + 18), (cx, y + layer_h + 34)], fill="#8391A2")
        y += layer_h + 32

    footer_y = 885
    d.rounded_rectangle((70, footer_y, 1730, 955), radius=18, fill="#17365D")
    draw_centered(d, (90, footer_y, 1710, 955), "边界：研究信号 ≠ 自动交易；模型输出 ≠ 事实；部署 ≠ 产品被采用；所有关键结论必须可追踪、可复核。", tiny_font, "#FFFFFF")
    img.save(path, dpi=(180, 180))


def add_toc(doc: Document):
    entries = [
        "0. 执行摘要：现在的 AR 到底变成了什么",
        "1. 最重要：未来产品落地计划与期望",
        "2. 第三周以来的重要变化：按时间与系统分类",
        "3. 当前研究框架与整体模型结构",
        "4. AIOS：机制、前后变化与真实开发进度",
        "5. 前后端与工程概况：做到哪里、还没做到哪里",
        "6. 当前开发总看板与关键卡点",
        "7. 人员调整与最新分工（6 人团队）",
        "8. 会议汇报建议稿（可直接照着讲）",
        "9. 建议向学长重点请教的四个问题",
        "附录 A. 证据索引与当前 PR 快照",
    ]
    for entry in entries:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.12)
        p.paragraph_format.space_after = Pt(2.5)
        add_text(p, entry, color=DEEP_BLUE, size=9.4)


def build() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    make_architecture_figure(FIG_PATH)
    doc = Document()
    configure_styles(doc)
    set_section(doc.sections[0])
    props = doc.core_properties
    props.title = "AR 第三周至今重要变化与下一阶段产品落地汇报"
    props.subject = "2026-07-20 至 2026-08-15 项目现状、架构、进度与团队汇报"
    props.author = "Junyan / AR"
    props.keywords = "AR, Research OS, AIOS, Product OS, roadmap, team"

    # Masthead / cover
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    add_text(p, "AR · MENTOR MEETING BRIEF", bold=True, color=BLUE, size=10)
    add_bottom_border(p)
    title = doc.add_paragraph(style="Title")
    title.paragraph_format.space_before = Pt(20)
    add_text(title, "AR 第三周至今重要变化\n与下一阶段产品落地汇报", bold=True, color=NAVY, size=24)
    sub = doc.add_paragraph(style="Subtitle")
    add_text(sub, "从研究项目走向“Research OS × AIOS × Product OS”的可验证产品系统", color=MUTED, size=12)

    meta = doc.add_table(rows=5, cols=2)
    set_table_width(meta, 9360)
    set_table_grid(meta, [1700, 7660])
    set_table_borders(meta, color="FFFFFF", size="0")
    labels = ["汇报人", "会议对象", "时间范围", "状态快照", "核心事实源"]
    values = [
        "Junyan（Founder / 总负责人）",
        "学长 / 高级审阅者",
        "PM 第三周 2026-07-20 — 2026-08-15",
        "GitHub main 截至 2026-08-15；最新主干提交 #269（5edc78e8）",
        "仓库主干、W30/W31 周报、Research/AI/Product 永久总账、当前 PR 快照",
    ]
    for i, (lab, val) in enumerate(zip(labels, values)):
        set_cell_width(meta.cell(i, 0), 1700)
        set_cell_width(meta.cell(i, 1), 7660)
        set_cell_shading(meta.cell(i, 0), LIGHT_GRAY)
        set_cell_margins(meta.cell(i, 0), 90, 120, 90, 120)
        set_cell_margins(meta.cell(i, 1), 90, 120, 90, 120)
        p1 = meta.cell(i, 0).paragraphs[0]; p1.paragraph_format.space_after = Pt(0)
        p2 = meta.cell(i, 1).paragraphs[0]; p2.paragraph_format.space_after = Pt(0)
        add_text(p1, lab, bold=True, color=NAVY, size=9.5)
        add_text(p2, val, size=9.5)

    doc.add_paragraph()
    add_callout(
        doc,
        "会议开场的一句话",
        "第三周以来，AR 最重要的变化不是“又增加了多少模型功能”，而是把研究、AI 工程和产品工程拆成三套相互约束的操作系统：Research OS 负责产生证据，AIOS 负责让 AI 工作可控且可证明，Product OS 负责把这些能力变成团队和未来用户真正能完成的工作。",
        fill=PALE_BLUE,
        accent=BLUE,
    )
    p = doc.add_paragraph(style="Small Note")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_text(p, "内部会议材料 · 状态词只描述工程成熟度，不构成投资结论或产品对外承诺。", italic=True)

    page_break(doc)
    add_heading(doc, "目录与阅读方式", 1)
    add_toc(doc)
    add_callout(
        doc,
        "建议会议顺序（15–20 分钟）",
        "先用第 1 节讲产品落地与未来期望（5 分钟）；再用第 3–5 节讲 Research OS、AIOS 与前后端真实进度（7 分钟）；最后用第 7 节讲人员调整（3 分钟），以第 9 节的四个问题向学长请教。",
        fill="F8FAFC",
        accent=DEEP_BLUE,
    )
    add_heading(doc, "状态口径", 2)
    add_status_table(
        doc,
        ["状态", "会议中应如何理解"],
        [
            ("已合并 / MAIN", "代码或文档已进入 main；仍不自动等于部署或被真实使用。"),
            ("验证中 / VALIDATING", "有实物并通过部分测试，但仍等待真实运行、集成或使用证明。"),
            ("PR 中 / IN REVIEW", "存在可审阅实现；不能描述为团队已拥有的生产能力。"),
            ("未接线 / DELIVERED_UNWIRED", "局部组件存在，但尚未接入唯一工作流或消费者。"),
            ("规划 / PROPOSED", "方向已定义，尚没有可运行交付。"),
            ("阻断 / BLOCKED", "关键依赖、数据或批准缺失；系统应显式失败而不是猜。"),
        ],
        [2100, 7260],
        status_col=0,
    )

    page_break(doc)
    add_heading(doc, "0. 执行摘要：现在的 AR 到底变成了什么", 1)
    add_callout(
        doc,
        "当前总判断",
        "AR 已经从“单人研究 + 若干模型脚本”进入“多人协作的研究工程系统”阶段。研究后端和治理基础进步最快；AIOS 已有可用地基但关键链路仍未闭合；产品前端仍在脚手架阶段。因此现在最优先的不是继续增加页面或 Agent，而是把已有能力闭合成一个可被团队每天使用、可由学长审阅、可用真实证据验收的内部产品。",
        fill=LIGHT_BLUE,
        accent=NAVY,
        status="关键结论",
    )
    add_heading(doc, "六个层面的当前状态", 2)
    add_status_table(
        doc,
        ["层面", "本阶段最重要变化", "当前状态", "关键缺口"],
        [
            ("产品方向", "明确内部研究工作台优先，外部产品延后到真实使用闭环之后", "战略已定义", "产品楔子、外部用户与付费/合作场景仍需验证"),
            ("Research OS", "统一研究宪法、全市场漏斗、宏观线、因果与随机对照开始成为正式制度", "主干 + 验证中", "U5 深度研究/决策闭环与前瞻性效果验证未完成"),
            ("Nightly / 后端", "研究漏斗 U1→U4 已以隔离观察模式接入夜链", "主干 / PARTIAL", "8 月 17 日首次真实定时验收；部分通道仍缺数据"),
            ("AIOS", "任务、适配器、能力路由、工作区、变异门进入主干", "地基已合并", "Policy、Context、Isolation 尚在 PR，未形成端到端 harness"),
            ("产品前端", "legacy 工具可用，新 web 主干已有静态 M0 PR", "PR 中 / 未接线", "无统一 contract client；新前端尚未连接真实 API 与核心旅程"),
            ("团队", "从 1 人扩展到 6 人，并按系统层分工学习 AI Engineering", "重新分工中", "Eric onboarding；RACI、Issue owner 与旧 PR 尚未完全同步"),
        ],
        [1200, 3250, 1500, 3410],
        status_col=2,
        font_size=8.7,
    )
    add_heading(doc, "从第三周到现在的量化变化", 2)
    add_status_table(
        doc,
        ["指标", "起点 / 历史节点", "截至 2026-08-15", "解释"],
        [
            ("组织", "W30 仍以 Junyan 单人为主", "6 人：Junyan、Simon、Reed、Better、Jason、Eric", "Junyan 退出 AIOS 日常开发，保留战略与高风险终审"),
            ("GitHub 交付", "W30 开始系统化 PR 流程", "07-20 至 08-15 共 76 个 PR 合并；当前 17 个 PR 仍开放", "交付速度高，但 PR 清账与所有权同步成为新瓶颈"),
            ("研究链", "L0.5 前兆、双层线、首次完整 paper-trade 复盘", "U0-U4 合同闭合并进入 nightly 观察", "从单票研究迭代到全市场、可复核的数据管线"),
            ("工程治理", "定时器重启导致结算缺失，质量门仍依赖人工", "事件账本、WAL、mutation gate、隔离发布、branch protection", "目标由“能跑”改为“能自证、失败可见”"),
            ("产品", "7631 行 legacy Dashboard 为主", "新 web M0 PR + Progress Board/只读 API", "产品主干尚未形成，当前主要还是内部工程工具"),
        ],
        [1450, 2600, 2800, 2510],
        font_size=8.6,
    )

    page_break(doc)
    add_heading(doc, "1. 最重要：未来产品落地计划与期望", 1)
    add_callout(
        doc,
        "建议对学长使用的产品定义",
        "AR 未来不是一个泛化聊天机器人，也不是自动交易工具；它首先是一套“决策级研究工作空间”：把市场信息、全市场筛选、宏观状态、组合风险、研究证据、AI 协作和事后复盘放进一个可追踪的工作流，让研究者在更短时间内获得更高质量、可验证、可审计的判断材料，最终决策仍由人完成。",
        fill=GREEN_BG,
        accent=GREEN,
    )
    architecture_shape = doc.add_picture(str(FIG_PATH), width=Inches(6.6))
    architecture_shape._inline.docPr.set(
        "descr",
        "AR 四层主链：工程与数据底座支撑 Research OS，AIOS Harness 管理任务与证据，产品体验层向用户提供驾驶舱、研究、组合、复盘与工作台。",
    )
    architecture_shape._inline.docPr.set("title", "AR 三层系统与产品落地主链")
    last = doc.paragraphs[-1]
    last.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = doc.add_paragraph(style="Small Note")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_text(cap, "图 1｜产品价值由 Research OS 的证据、AIOS 的可控执行与 Product OS 的用户工作流共同产生。")

    add_heading(doc, "1.1 产品楔子：先服务内部，再验证外部", 2)
    add_status_table(
        doc,
        ["阶段", "优先用户", "必须完成的产品", "退出门 / 成功标准"],
        [
            ("P0 内部 Alpha\n未来 0–4 周", "Junyan + 团队 + 学长", "唯一 canonical app；今日驾驶舱；研究漏斗；组合/复盘；/team 与 /health；审阅包导出", "核心数字 parity；freshness/blocked 可见；团队连续 2 周完成日常研究与周审阅"),
            ("P1 团队 Beta\n未来 4–8 周", "固定审阅者 + 研究协作者", "U0-U5 可下钻；Macro→行业→组合；AIOS 运行证据；反馈与 Memory 闭环", "六条旅程无 P0/P1；失败可定位到 owner；部署可回滚；使用时间明显下降"),
            ("P2 封闭试点\n未来 8–12+ 周", "3–5 位目标研究者/合作方", "权限、隐私、成本、支持、审计、产品叙事与 onboarding", "真实用户重复使用并愿意继续；价值与边界能被非团队成员理解；无自动交易能力"),
            ("P3 对外准备\n门槛触发后", "已验证细分用户", "商业模式、SLA、合规、数据许可、外部安全与运维", "由 Junyan 单独批准；不能因为 UI 完成就提前宣称成熟"),
        ],
        [1700, 1650, 3410, 2600],
        font_size=8.3,
    )

    add_heading(doc, "1.2 未来产品的七个核心页面与三个工具面", 2)
    add_status_table(
        doc,
        ["产品面", "用户要完成的工作", "最小证据要求"],
        [
            ("今日驾驶舱", "3 分钟完成市场、宏观、组合风险与今日待核对事项", "as_of、来源、状态、owner；不把 PARTIAL 显示为 COMPLETE"),
            ("市场研究", "从 U0 全市场进入 U1 扫描、U2 候选、U3 电池、U4 人审、U5 深研", "每次晋级 reason code、随机对照、不可追溯样本不得消失"),
            ("宏观", "从事件下钻到状态、行业与组合暴露", "官方数据/日期/修订；只改变优先级和风险预算，不直接生成交易动作"),
            ("模型组合", "查看 NAV、现金、持仓、风险、数据时间和 paper/human 区分", "账本/契约逐字段一致，来源可回溯"),
            ("交易复盘", "按周查看 reasoning、结果、违约、反思和下一步", "两票、双层线、registered→filled→exit 全链"),
            ("周报与 Memory", "让学长不读代码也能理解本周变化并给可追踪建议", "事实、判断、未决问题、owner、下一检查点"),
            ("研究工作台", "单票图表、事件、factpack、法庭与证据核对", "前端只展示批准契约，不重算研究结论"),
            ("/team /health /settings", "看任务、AIOS、数据源、nightly、报警、部署；管理非敏感偏好", "只读投影、无浏览器 secret、报警必须有 owner/ack/status"),
        ],
        [1600, 4050, 3710],
        font_size=8.3,
    )

    add_heading(doc, "1.3 我们对产品落地的期望", 2)
    for text, prefix in [
        ("效率：把盘前检查、候选进入深研、周度复盘和会议汇报从“拼文件”变成固定路径。", "效率："),
        ("质量：任何数字都能看到日期、来源、完整性与失败状态；AI 结论必须绑定证据。", "质量："),
        ("协作：任务、上下文、PR、证据与 Memory 是一个闭环，换设备或换模型不再丢失工作。", "协作："),
        ("学习：五位 AIOS 建设成员各自交付真实系统层，形成可复述、可测试、可展示的 AI Engineering 作品。", "学习："),
        ("外部价值：先证明固定研究者愿意持续使用，再讨论规模化、收费或合作；拒绝用“页面上线”替代用户价值。", "外部价值："),
    ]:
        add_bullet(doc, text, bold_prefix=prefix)

    add_heading(doc, "1.4 当前不做什么", 2)
    for item in (
        "不把 AR 宣称为自动交易、投顾或能够自主下单的系统。",
        "不在核心旅程未闭合前继续横向增加第八、第九个页面。",
        "不把模型输出直接当作事实、批准或研究结论。",
        "不声称拥有 at-scale 的高并发、跨云或成熟 SLA；当前重点是 bounded internal system。",
        "不在权限、隐私、数据许可、成本和支持模型未验证前对外开放。",
    ):
        add_bullet(doc, item)

    page_break(doc)
    add_heading(doc, "2. 第三周以来的重要变化：按时间与系统分类", 1)
    add_status_table(
        doc,
        ["阶段", "战略变化", "研究 / 模型", "工程 / 产品 / 团队"],
        [
            ("W30\n07-20—07-24", "开始从“看结果”转向“记录规则与违约”", "L0.5 前兆层；双层线与破线复核；首个 registered→filled→target exit 闭环；事实包形成", "第五次结算漏跑暴露单机定时/重启风险；开始把故障记入制度"),
            ("W31\n07-27—08-01", "主题改为“从能跑到能自证”", "红旗闸门、六维电池、10 天检察官、方向分列归因、Macro 研究线", "团队 1→3；Nightly v2 15 步；public/data/v2 契约；主分支保护；敌意审计与 mutation 思路"),
            ("W32\n08-02—08-08", "建立 Research OS × AI OS × Product OS 三套永久地基", "全市场 U0、E1 事件层、point-in-time feature store、宏观官方数据与 MRG", "统一 AgentAdapter / Kimi / DeepSeek；Progress API；数据发布改为可恢复与可验证"),
            ("W33\n08-09—08-15", "从模块建设转入闭环、产品化和多人分层", "Macro M1-C 接入 nightly；U1-U4 合同闭合；研究漏斗进入隔离观察；R-035 评估启动", "AIOS task registry/router/workspace/mutation attribution 入 main；新 web M0 PR；团队扩至 6 人、Eric 加入"),
        ],
        [1350, 2200, 2900, 2910],
        font_size=8.35,
    )

    add_heading(doc, "2.1 最重要的五个“性质变化”", 2)
    items = [
        ("从单点模型到系统架构", "模型只是 Research OS 的一个组件；数据、合同、nightly、账本、审阅和产品都成为一等公民。"),
        ("从自述完成到证据完成", "引入 fail-closed、mutation tests、独立 review、append-only 事件与 hash 绑定，完成不能由执行者自签。"),
        ("从少量关注票到全市场漏斗", "U0-U5 让全市场扫描、候选晋级、随机对照、六维电池、人审和后续验证成为可重放链路。"),
        ("从本地个人工作到 GitHub 协作", "分支保护、Issue/PR、worktree、同步 doctor 和 Progress Board 让团队可在 Windows/Mac 多端工作。"),
        ("从研究展示到产品工作流", "产品目标不再是堆图表，而是让 Junyan、审阅者和建设者完成固定任务，并能看到失败与责任人。"),
    ]
    for title_text, body in items:
        p = doc.add_paragraph(style="List Number")
        add_text(p, title_text + "：", bold=True, color=DEEP_BLUE)
        add_text(p, body)

    add_heading(doc, "2.2 W31 研究归因给我们的关键纠偏", 2)
    add_callout(
        doc,
        "历史审计结果（不是未来收益承诺）",
        "W31 对 39 条方向判断的回溯显示：系统对“危险/谨慎”更擅长，对“进攻/建设性观点”尚未证明。建设性样本 n=17，原始命中约 18%、超额命中约 29%；谨慎样本 n=22，原始命中约 64%、超额命中约 55%。因此未来评测必须永久分开 offense 与 defense，不能用一个总准确率掩盖弱项。",
        fill=AMBER_BG,
        accent=AMBER,
    )

    page_break(doc)
    add_heading(doc, "3. 当前研究框架与整体模型结构", 1)
    add_callout(
        doc,
        "Research OS 的机制",
        "Research OS 不追求一次性给出“答案”，而是把研究拆成不可跳步的证据链：先定义覆盖范围与时间点，再用独立通道发现候选，用随机对照和六维电池检验，再由人类晋级，最后通过 factpack、决策书、后续价格与因果簇验证，让好观点与坏观点都能被追溯。",
        fill=GREEN_BG,
        accent=GREEN,
    )
    add_heading(doc, "3.1 研究主链", 2)
    add_status_table(
        doc,
        ["层", "职责", "当前进度", "下一道门"],
        [
            ("U0 Universe", "全市场证券注册、资格与 point-in-time 身份", "已合并 / MAIN", "持续更新与来源完整性"),
            ("U1 Scan", "价格量、事件、基本面、资金、宏观等独立通道扫描", "已合并；nightly 观察", "解决 FUNDAMENTAL/FUND_FLOW/MACRO 大面积缺数"),
            ("U2 Candidate", "候选、reason code、主通道与同批随机对照", "合同已合并", "真实运行积累足够批次"),
            ("U3 Battery", "六维完整性与红旗，不允许挑维度", "合同已合并", "与真实 factpack/decision sheet 进一步绑定"),
            ("U4 Human Review", "人审 packet、3–5 名晋级、身份与 receipt 边界", "合同已合并；离线闭环 PR #272", "把人审与生产链安全连接"),
            ("U5 Deep Research", "factpack、wrong-if、决策书、法庭与结果追踪", "DATA_BLOCKED / 部分能力", "完成统一闭环，并在前瞻数据上验证"),
            ("R-035 Evaluation", "T+1/T+3/T+5/T+10 离线评估，分主通道/对照与 battery pass", "PR #274", "依赖 #273 生产验收，随后再谈生产接线"),
        ],
        [1250, 3550, 1950, 2610],
        status_col=2,
        font_size=8.45,
    )

    add_heading(doc, "3.2 研究系统外围结构", 2)
    add_status_table(
        doc,
        ["模块", "作用", "当前真实状态"],
        [
            ("Unified Research OS v1.5", "研究宪法：证据层级、两票、双层线、红旗、六维、方向分列、宏观边界", "主干权威；任何改变需版本与 Junyan 批准"),
            ("Macro OS", "官方数据采集、双状态/MRG、行业与组合消费者", "M0-A/B/B2/B3 与 M1-A/B/C 已合并；在 nightly 中隔离校准"),
            ("Nightly v4", "结算、数据、研究、Macro、发布与报警的唯一夜间编排", "主干运行；新漏斗隔离、PARTIAL 上浮；真实定时验收待 #273"),
            ("证据/账本", "append-only ledger、manifest、WAL、hash、迁移与可恢复发布", "主干；避免旧文件冒充和迁移不可逆"),
            ("数据合同", "public/data/v2、schema、producer/consumer/freshness/data_status", "主干持续增长；产品消费者仍未统一"),
        ],
        [1900, 4150, 3310],
        font_size=8.55,
    )

    add_heading(doc, "3.3 研究框架目前卡在哪里", 2)
    blockers = [
        "第一次真实 launchd 定时验收尚未发生：#273 明确以 2026-08-17 16:35 的真实运行作为最终验收，当前只能称“手工/离线证明充分、生产定时待证”。",
        "漏斗数据质量目前为 PARTIAL：观察结果明确暴露 FUNDAMENTAL_VALUATION、FUND_FLOW_CHIPS、MACRO_CROSS_ASSET 等降级通道；不能把编排 COMPLETE 说成数据 COMPLETE。",
        "U5 尚未闭合：factpack、decision sheet、wrong-if、前瞻因果簇与结果评价还没有被同一链路证明。",
        "进攻能力未证明：当前回溯更支持风险识别能力；需要前瞻、分方向、与随机对照比较的样本。",
        "对外产品不能把研究信号写成交易动作：Macro 与漏斗在当前阶段必须保持观察、non-blocking、claim_allowed=false。",
    ]
    for b in blockers:
        add_bullet(doc, b)

    page_break(doc)
    add_heading(doc, "4. AIOS：机制、前后变化与真实开发进度", 1)
    add_callout(
        doc,
        "AIOS 的一句话机制",
        "AIOS 是套在模型外部的工程 Harness。它先把需求编译成任务合同，加载最少且最新的权威上下文，再按权限、风险、预算和能力选择 Agent，在隔离环境执行；执行结果必须经过独立验证、敌意评测和人类门，最后把证据、成本、失败与 Memory 写回。目标是让 AI 的工作可控、可复现、可审计、可协作，而不是让 AI 拥有更大的自由。",
        fill=AMBER_BG,
        accent=AMBER,
    )
    add_heading(doc, "4.1 AIOS 端到端工作流", 2)
    flow = [
        "需求进入：Issue / 产品需求 / 研究任务 → ai-task 合同；模糊、缺范围或缺验收则 SPEC_BLOCKED。",
        "权威解析：Constitution Graph + Repository Intelligence 找出规则、代码、数据生产者/消费者和冲突。",
        "上下文组装：Context Broker 只加载该阶段需要的最新片段，记录 hash、版本、freshness 与外部不可信输入。",
        "决策与路由：Policy Gate 检查文件、网络、预算、风险；Capability Router 只选择通过相应评测的 Agent/Skill。",
        "隔离执行：独立分支/worktree、最小工具权限、timeout、幂等、运行 manifest；不得直接写 main。",
        "验证与审核：diff-aware tests、contract tests、mutation tests、secret/scope checks、独立 adversarial review。",
        "人类门：核心/高风险/宪法变更由 Junyan 最终审核；日常低风险工作不需要 Junyan 逐步授权。",
        "回写学习：Task/PR/证据/成本/报警/Memory/Backlog 同步，避免下一次从聊天重新开始。",
    ]
    for item in flow:
        add_number(doc, item)

    add_heading(doc, "4.2 十二个 Harness 模块的状态", 2)
    add_status_table(
        doc,
        ["模块", "功能", "当前状态", "关键证据 / 缺口"],
        [
            ("H1 Runtime", "统一运行、retry、timeout、幂等与生命周期", "部分", "Adapter 有；真实 executor/sandbox 尚未闭合"),
            ("H2 Constitution", "宪法、权威顺序、核心文件与修订门", "设计", "V3 蓝图本地候选；Constitution Graph 待 Simon"),
            ("H3 Task Lifecycle", "Task contract、registry、state、reconciler", "主干地基", "#248 已合并；端到端 reconciliation 仍需扩展"),
            ("H4 Workflow Twin", "将团队实际工作流变成可追踪状态机", "部分", "Progress Board 有；Research/Product/AIOS 统一投影待建"),
            ("H5 Repo Intelligence", "跨文件依赖、owner、change map、影响分析", "规划", "Reed 下一主切片"),
            ("H6 Context Broker", "最少/最新/可追踪上下文与 RAG", "PR 中", "#265 A-010 Context Builder；真实检索与 Broker 未接线"),
            ("H7 Prompt / Role / Skill", "按 case 生成角色、格式、工具与完成标准", "规划", "需 Better 将 prompt engineering 变成任务类型产品层"),
            ("H8 Agent / Skill Router", "provider 适配、能力注册、任务路由", "主干地基", "#230/#232/#237 adapters + #246 router；尚未接 Policy/Executor"),
            ("H9 Evals / Quality", "评测、对抗、mutation、质量回归", "部分主干", "mutation attribution #268；统一 Harness eval 矩阵待 Jason"),
            ("H10 Observability", "run、成本、延迟、失败、健康与报警", "部分", "nightly/ledger 有；AIOS 统一 telemetry 待 Eric"),
            ("H11 Safety / Authority", "Policy、scope、secrets、human gate", "PR 中", "#250 Policy、#261 Isolation Planner；独立审批链未闭合"),
            ("H12 Memory / Learning", "任务复盘、根因簇、周 digest、技能演化", "规划", "Backlog 已定义，尚无端到端真实任务证明"),
        ],
        [1400, 2700, 1500, 3760],
        status_col=2,
        font_size=8.05,
    )

    add_heading(doc, "4.3 已进入 main 与仍在 PR 的分界", 2)
    add_status_table(
        doc,
        ["状态", "内容", "会议口径"],
        [
            ("已合并", "AgentAdapter #230、Kimi #232、DeepSeek #237、Task Registry #248、Capability Router #246、治理 mutation #253/#255/#268、Team AI Workspace #260", "可以说“地基存在并受离线测试约束”；不能说“AIOS 已端到端上线”"),
            ("当前 AIOS PR", "#250 Policy、#265 Context Builder、#261 Isolation Planner、#266 Workspace Sync，以及多个旧 seed/parked PR", "需要 rebase、独立 Review、合并/关闭决策；否则新成员会读到多套真相"),
            ("蓝图状态", "AIOS 总蓝图 v3 在本地 worktree 已形成统一候选", "尚未进入 GitHub main；必须开 PR 后才能成为团队权威"),
            ("开放 PR 总量", "当前 17 个：AIOS 10、研究/数据 5、产品/前端 2", "交付已从“写代码”转为“清理依赖、证明状态和闭合主线”"),
        ],
        [1500, 4800, 3060],
        status_col=0,
        font_size=8.55,
    )

    page_break(doc)
    add_heading(doc, "5. 前后端与工程概况：做到哪里、还没做到哪里", 1)
    add_heading(doc, "5.1 前端", 2)
    add_status_table(
        doc,
        ["资产", "当前实物", "状态", "真实结论"],
        [
            ("Legacy Dashboard", "src/Dashboard.jsx，约 7631 行", "LIVE / LEGACY", "可演示、可维护 P0，但不可继续作为新产品主干"),
            ("Team Progress", "#175 只读 API、#225 endpoint 对齐、#238 实时蓝图面板", "已合并", "内部协作可见性增强，但仍在 legacy 载体"),
            ("新 web 主干", "#257 React 18 + Vite + TypeScript、主题 tokens、七个导航 shell", "PR 中", "build/lint/响应式已验证；没有生产数据、API、写路径或模型调用"),
            ("Contract Client", "runtime validator、typed client、data-state 七态", "未接线", "Issue #198 / PE-M2 尚未形成产品消费者；这是产品化核心瓶颈"),
            ("核心旅程", "驾驶舱、市场、Macro、组合、复盘、周报、工作台", "新主干未完成", "当前只有导航/工具和 legacy 片段；不能说前端产品已完成"),
            ("部署", "GitHub Pages + Vercel 双路径", "验证中", "canonical URL、API host、preview/prod、rollback 仍需 M0 ADR"),
        ],
        [1700, 3400, 1550, 2710],
        status_col=2,
        font_size=8.35,
    )

    add_heading(doc, "5.2 后端与数据", 2)
    add_status_table(
        doc,
        ["层", "已完成 / 在主干", "仍需完成"],
        [
            ("Research backend", "Nightly v4、U0、E1、point-in-time feature store、U1-U4 funnel、Macro M1-C、证据/发布校验", "U5、R-035、真实 launchd acceptance、缺数通道修复"),
            ("Contracts / API", "public/data/v2、多个 schema 与 team-progress read API", "统一 producer→schema→BFF→client；清理 api/ 中 read/model 混合与旧静态上下文"),
            ("AIOS backend", "Adapter、Task Registry、Router、Workspace、mutation gates", "Policy + Context + Executor + Evals + Telemetry + Memory 端到端接线"),
            ("Operations", "branch protection、PR review、WAL、可恢复迁移、worktree 隔离", "跨成员同步 doctor 常态化；Windows/Mac 一致性；open PR 清账"),
            ("Production acceptance", "手工 canary、离线/对抗测试证据充分", "2026-08-17 真实定时验收；连续运行窗口与报警消费"),
        ],
        [1800, 4300, 3260],
        font_size=8.55,
    )

    add_callout(
        doc,
        "需要避免的误解",
        "后端研究能力比前端成熟得多，但“有很多 JSON / API / 夜链步骤”不等于产品已经可用。产品成熟要求同一个真实用户旅程贯穿数据生产、契约、BFF、前端、错误状态、部署和反馈；目前这条线尚未闭合。",
        fill=RED_BG,
        accent=RED,
    )

    add_heading(doc, "5.3 工程治理的升级", 2)
    for item in (
        "GitHub 成为唯一共享事实源：main 受保护，日常开发走 Issue → branch/worktree → PR → CI → review → merge。",
        "多端迁移由 Team AI Workspace v1 约束：优先使用现有 clone、记录 origin/main SHA、doctor 检查、生产数据与本地密钥不入仓库。",
        "新组件先观察再获得权力：研究漏斗进入 nightly 时为 isolated/non-blocking，证明稳定后再申请正式阻断权。",
        "测试从“验证能成功”升级到“证明关键门删掉会失败”：mutation tests 与独立对抗复核成为治理门。",
        "发布与迁移变得可恢复：manifest、WAL、hash、complete stamp 和 rollback 顺序共同避免两头落空。",
    ):
        add_bullet(doc, item)

    page_break(doc)
    add_heading(doc, "6. 当前开发总看板与关键卡点", 1)
    add_status_table(
        doc,
        ["优先级", "卡点", "为什么重要", "解除条件 / Owner"],
        [
            ("P0", "产品主干没有端到端真实旅程", "前端仍是 shell/legacy，无法证明产品价值", "Better：#257 处置 + M0 ADR + PE-M2 contract client；Junyan 批 canonical app"),
            ("P0", "Nightly 真实定时验收未完成", "手工成功不能替代系统按时、持续、可报警运行", "#273；2026-08-17 16:35 真实 launchd acceptance"),
            ("P0", "AIOS Policy→Context→Router→Executor 未闭合", "模型仍不能在统一权限与上下文下安全执行", "Simon/Reed/Jason/Eric：#250/#265/#261 rebase、review、接线"),
            ("P1", "U5 与前瞻评估未闭合", "有候选不等于研究产出质量被证明", "Research：#272/#274；factpack/decision/wrong-if/因果簇"),
            ("P1", "17 个 open PR 与旧 seed 并存", "多套协议会导致新成员加载陈旧上下文", "Simon 组织清账；owner 给出 merge/close/supersede 终态"),
            ("P1", "团队 RACI 与 backlog 未完全同步", "Eric 新加入、Junyan 退出 AIOS 开发后，旧 owner 字段已过时", "Simon 48 小时内更新 Issue、Backlog、Review pair 与 DoD"),
            ("P2", "数据通道 PARTIAL", "漏斗有结构但部分维度覆盖不足", "研究/数据 owner 分通道修复，保持显式 degraded_channels"),
        ],
        [800, 2600, 2750, 3210],
        font_size=8.3,
    )

    add_heading(doc, "6.1 未来四周建议执行顺序", 2)
    add_status_table(
        doc,
        ["周", "产品主线", "Research / AIOS 主线", "退出证据"],
        [
            ("Week 1", "锁 Product Brief、M0 ADR、#257 merge/close；确定一个 canonical app", "完成 #273 真实验收；AIOS V3 开 PR；17 PR 分类清账", "唯一 URL/事实源；真实定时 receipt；每个 PR 有 disposition"),
            ("Week 2", "PE-M2 contract client + 七种 data state；先接 /team /health", "合并或关闭 #250/#265/#261/#266；Eric 建 run/telemetry contract", "错误/STALE/PARTIAL 可复现；AIOS 最小 run 有 manifest 与证据"),
            ("Week 3", "接一条完整用户旅程：盘前驾驶舱或全市场研究", "Repo Intelligence + Context Broker 最小切片；Harness eval 回归矩阵", "从源数据到页面到反馈全链可演示"),
            ("Week 4", "内部 Alpha 连续使用、周报导出、回滚演练", "端到端 Policy→Context→Router→Executor→Verify 低风险任务", "连续运行记录；失败有 owner；Junyan 只在终审门出现"),
        ],
        [1000, 3100, 3350, 1910],
        font_size=8.25,
    )

    page_break(doc)
    add_heading(doc, "7. 人员调整与最新分工（6 人团队）", 1)
    add_callout(
        doc,
        "权限与责任原则",
        "Junyan 不参与 AIOS 日常开发，不成为每一步的瓶颈；所有成员可在授权分支、Issue 和 PR 中正常建设。Junyan 只保留产品方向、研究宪法、main/核心文件、高风险权限与最终上线的总览审核。任何成员都不能单独 merge 自己负责的关键层，必须经过独立 cross-review。",
        fill=LIGHT_BLUE,
        accent=NAVY,
    )
    add_status_table(
        doc,
        ["成员", "当前角色", "主要拥有的系统层", "未来 2 周首个可验收交付", "独立 Reviewer"],
        [
            ("Junyan", "Founder / 总负责人 / Product & Research Authority", "产品楔子、研究宪法、main/core、高风险终审；不写 AIOS", "批准 Product Brief + canonical app ADR；定义外部试点进入门；审查核心合并", "学长/指定技术 reviewer 提建议；Junyan 最终裁决"),
            ("Simon", "项目主管 / AI Technical PM / Control Plane Architect", "H2 Constitution、H3 Task、H4 Workflow Twin；团队节奏", "把 V3、backlog、17 PR、六人 RACI 合成一张实时 roadmap；交 Constitution Graph 最小切片", "Better（流程/UX）+ Jason（authority）"),
            ("Reed", "Agent Platform / Repo Intelligence / Skills", "H5 Repo Intelligence、H8 Router/Agent/Skill；H6 工程接口", "完成 change-map / dependency-map；推进 #265；把 Router 与上下游合同对齐", "Jason（安全/测试）"),
            ("Better", "Product Engineer / Context & AI Experience", "Product OS、web 前端、BFF/contract client、H6/H7 用户层", "决定 #257；完成 M0 ADR；用 fixture 实现七态 client；Context Pack/RAG 最小 demo", "Simon（架构）+ Jason（P5 安全）"),
            ("Jason", "Evals / Safety / Governance Engineer", "H9 Evals、H11 Safety/Human Gate；对抗 Review", "完成 Harness Eval 最小矩阵；独立审查 #250/#261/#265；建立不能自签的 release gate", "Reed（测试实现）+ Simon（状态语义）"),
            ("Eric", "Runtime / LLMOps / Observability Engineer", "H1 Runtime、H10 Observability、H12 运行 Memory 接口", "完成 workspace onboarding；定义 ai-run/health/cost/alert contract；做一个离线 executor+telemetry vertical slice", "Jason（权限/失败）+ Reed（runtime 接口）"),
        ],
        [900, 1800, 2050, 3350, 1260],
        font_size=7.8,
    )

    add_heading(doc, "7.1 共同技能底座", 2)
    for item in (
        "Git/GitHub：Issue、branch/worktree、rebase、PR、CI、review、merge/close、recover。",
        "Contract-first：schema、fixture、producer/consumer、freshness、data_status 与 backward compatibility。",
        "Evaluation-first：先写失败样例、对抗样例和 mutation，再相信“测试全绿”。",
        "AI Harness：context、prompt、tool、policy、runtime、evidence、memory 是一个系统，不是单独学 prompt engineering。",
        "可观测与安全：run_id、hash、cost、latency、owner、alert、secret、scope、network、human gate。",
        "产品思维：从需求和用户工作出发；每个功能说明谁用、完成什么、如何失败、何时停止投入。",
        "复盘表达：每个 PR 必须能用“问题—机制—证据—残余风险—下一步”五句话讲清楚。",
    ):
        add_bullet(doc, item)

    add_heading(doc, "7.2 团队稳定工作流", 2)
    add_status_table(
        doc,
        ["步骤", "唯一记录位置", "完成条件"],
        [
            ("1. 需求/问题", "GitHub Issue + 对应 Research/AI/Product Backlog ID", "owner、范围、非目标、依赖、风险、验收、reviewer 完整"),
            ("2. 认领", "Issue comment / Progress Board CLAIM", "记录 branch、worktree、base SHA、ETA；避免重叠写"),
            ("3. 开发", "命名 branch/worktree", "只改 scope；本地/机密/生产状态不入仓库"),
            ("4. 自检", "PR evidence", "正向 + 负向 + 变异/对抗；记录真实命令和 gap"),
            ("5. 独立 Review", "PR review thread", "reviewer 不得是执行者；对 exact head；高风险需 Jason/Junyan"),
            ("6. 合并", "受保护 main", "CI 与 review 对当前 head 有效；旧 approval 不跨 commit"),
            ("7. 部署/验收", "release manifest / health / user receipt", "上线、回滚、数据 parity、真实使用分开记录"),
            ("8. Memory", "Backlog + weekly digest", "根因、残余风险、新债、owner、下一检查点回写"),
        ],
        [1150, 3000, 5210],
        font_size=8.45,
    )

    page_break(doc)
    add_heading(doc, "8. 会议汇报建议稿（可直接照着讲）", 1)
    add_heading(doc, "8.1 90 秒开场", 2)
    p = doc.add_paragraph()
    add_text(p, "“从第三周到现在，AR 最大的改变不是多做了几个功能，而是完成了三次转型。第一，从单人研究脚本转向 Research OS，把证据、全市场漏斗、宏观、账本和评估放进统一制度；第二，从模型能输出转向 AIOS，让任务、上下文、权限、执行和验证可追踪；第三，从展示研究结果转向 Product OS，目标是一个研究者每天真正会使用的决策级工作空间。现在研究后端最成熟，AIOS 地基已经进入主干，但关键 Harness 还没闭环；新前端仍是 M0 脚手架。因此下一阶段我们不是继续堆功能，而是先闭合一个内部产品旅程，再用真实使用证明是否值得对外。”")

    add_heading(doc, "8.2 建议重点讲的四页", 2)
    add_number(doc, "产品落地：内部 Alpha → 团队 Beta → 封闭试点；强调不是自动交易，不提前承诺规模化。")
    add_number(doc, "Research OS：U0-U4 已闭合并进入 nightly 观察；U5、真实定时验收与前瞻评估仍是关键门。")
    add_number(doc, "AIOS / 前后端：主干已有 task/router/workspace/mutation；policy/context/executor 在 PR；新 web 只有静态 shell。")
    add_number(doc, "团队：六人按层分工；Junyan 从 AIOS 日常开发退出，Simon 管项目，五位成员用建设层作为学习路径。")

    add_heading(doc, "8.3 如果学长追问“你们现在最缺什么”", 2)
    add_callout(
        doc,
        "建议回答",
        "我们现在不缺想法和模块，最缺的是“端到端收口”：一个真实用户旅程、一条真实定时运行、一个完整 AIOS 低风险任务，以及一套不会因成员增加而分叉的权威工作流。未来四周的成功不是 PR 数继续增长，而是这四条线都能被真实使用和独立复核。",
        fill=AMBER_BG,
        accent=AMBER,
    )

    add_heading(doc, "8.4 对容易被质疑的问题先主动说明", 2)
    add_status_table(
        doc,
        ["质疑", "建议答法"],
        [
            ("“是不是做得太大？”", "是，因此按内部产品楔子收缩；先闭合盘前/全市场研究中的一条，不同时做所有外部能力。"),
            ("“研究表现证明了吗？”", "没有。当前只证明了流程和风险识别的一部分；进攻能力、样本量与前瞻验证仍不足。"),
            ("“AIOS 是不是重复造平台？”", "不是通用平台；只实现 AR 工作流需要的 task/context/policy/eval/evidence/memory，优先复用 GitHub 与现有模型。"),
            ("“前端到哪了？”", "legacy 可用，内部 Progress 已有；新 web M0 静态脚手架在 PR，未接真实数据，真正产品旅程尚未完成。"),
            ("“六个人会不会管理失控？”", "风险存在；通过一人主层、独立 cross-review、统一 Issue/Backlog、Junyan 只守高风险门降低协调成本。"),
        ],
        [2500, 6860],
        font_size=8.7,
    )

    page_break(doc)
    add_heading(doc, "9. 建议向学长重点请教的四个问题", 1)
    questions = [
        ("产品楔子", "在“盘前驾驶舱 / 全市场研究 / 周度复盘”三条内部旅程中，哪一条最适合作为第一个可对外解释、可重复使用的产品楔子？为什么？"),
        ("研究可信度", "在 defense 明显优于 offense、样本仍小的阶段，学长认为最合理的前瞻验证周期、对照组与停止条件是什么？"),
        ("外部试点门", "从内部工具进入 3–5 人封闭试点前，除了功能完整，还必须达到哪些可靠性、隐私、数据许可或组织门槛？"),
        ("团队学习与交付", "五位成员一边学 AI Engineering 一边建设 AIOS，怎样设置阶段性产出，既能形成真实能力，又不让学习任务拖慢产品闭环？"),
    ]
    for idx, (title_text, body) in enumerate(questions, 1):
        add_callout(doc, f"Q{idx} · {title_text}", body, fill="F8FAFC", accent=DEEP_BLUE)

    add_heading(doc, "9.1 会后希望形成的决定", 2)
    for item in (
        "批准或修正 AR 的第一产品楔子与三阶段落地顺序。",
        "确认 Research OS 的前瞻验证标准，以及 offense/defense 是否继续分列。",
        "确认未来四周的三条硬门：真实 nightly、AIOS 最小闭环、产品单旅程。",
        "确认团队分工与 cross-review 结构是否足以支撑六人协作。",
        "把学长建议写成 Issue/Decision，而不是只停留在会议聊天。",
    ):
        add_bullet(doc, item)

    page_break(doc)
    add_heading(doc, "附录 A｜证据索引与当前 PR 快照", 1)
    add_heading(doc, "A.1 仓库事实源", 2)
    sources = [
        ("最新 main", "5edc78e8 — feat(nightly): wire research funnel as an isolated observation step (#269)"),
        ("周报", "docs/team/weekly/2026-W30.md；docs/team/weekly/2026-W31.md"),
        ("研究宪法", "docs/research/UNIFIED_RESEARCH_OS.md；ALL_MARKET_RESEARCH_FUNNEL.md；RESEARCH_ENGINEERING_BACKLOG.md"),
        ("AIOS", "docs/llm/AI_OS_ENGINEERING_BACKLOG.md；本地候选 docs/llm/AR_AIOS_MASTER_BLUEPRINT_v3.md"),
        ("产品", "docs/product/PRODUCT_ENGINEERING_BUILD_GUIDE.md；PRODUCT_ENGINEERING_BACKLOG.md"),
        ("团队", "docs/team/TEAM_CHARTER_v3.md；TEAM_AI_WORKSPACE_V1.md；PM_OPERATING_SYSTEM.md"),
    ]
    add_status_table(doc, ["类别", "权威证据"], sources, [1700, 7660], font_size=8.7)

    add_heading(doc, "A.2 截至 2026-08-15 的开放 PR 分类", 2)
    add_status_table(
        doc,
        ["分类", "数量", "PR", "会议意义"],
        [
            ("Research / Data", "5", "#239 #263 #272 #273 #274", "数据发布、U1-U5 离线闭环、真实定时验收、R-035 评估"),
            ("AIOS", "10", "#167 #215 #216 #218 #219 #250 #261 #262 #265 #266", "四个当前关键 PR 与六个旧 seed/parked/superseded 债需要收口"),
            ("Product / Frontend", "2", "#174 #257", "Progress legacy handoff 与新 web M0 scaffold 的产品裁决"),
        ],
        [1800, 700, 3400, 3460],
        font_size=8.6,
    )

    add_heading(doc, "A.3 近期关键 GitHub 证据", 2)
    links = [
        ("#269 研究漏斗接入 nightly 观察模式", "https://github.com/Lateily/Alpha-Research/pull/269"),
        ("#267 U1-U4 funnel contracts", "https://github.com/Lateily/Alpha-Research/pull/267"),
        ("#273 scheduled production acceptance audit", "https://github.com/Lateily/Alpha-Research/pull/273"),
        ("#274 R-035 funnel evaluation", "https://github.com/Lateily/Alpha-Research/pull/274"),
        ("#246 AIOS capability registry/router", "https://github.com/Lateily/Alpha-Research/pull/246"),
        ("#265 AIOS Context Builder", "https://github.com/Lateily/Alpha-Research/pull/265"),
        ("#257 M0 frontend scaffold", "https://github.com/Lateily/Alpha-Research/pull/257"),
        ("#260 Team AI Workspace", "https://github.com/Lateily/Alpha-Research/pull/260"),
    ]
    for label, url in links:
        p = doc.add_paragraph(style="List Bullet")
        add_hyperlink(p, label, url)

    add_heading(doc, "A.4 数据与表述边界", 2)
    p = doc.add_paragraph(style="Small Note")
    add_text(
        p,
        "本材料使用 2026-08-15 的 GitHub main 与开放 PR 快照。W30/W31 的组合或命中数字只作为历史系统审计背景，不代表当前净值，也不构成收益承诺。PR 描述中的测试通过只证明对应代码头部与范围；未合并、未部署、未真实使用的能力分别标为 PR 中、验证中或未接线。",
    )

    doc.save(OUT_PATH)
    print(OUT_PATH)


if __name__ == "__main__":
    build()
