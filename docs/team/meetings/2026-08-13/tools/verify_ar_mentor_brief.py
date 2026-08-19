from pathlib import Path
import zipfile

from docx import Document


p = Path(r"C:\Users\jctx\Desktop\AR\output\meetings\2026-08-15_AR_第三周至今_学长会议系统汇报.docx")
d = Document(p)
body_text = "\n".join(x.text for x in d.paragraphs)
table_text = "\n".join(
    cell.text
    for table in d.tables
    for row in table.rows
    for cell in row.cells
)
all_text = body_text + "\n" + table_text
headings = [x.text for x in d.paragraphs if x.style.name.startswith("Heading")]
with zipfile.ZipFile(p) as zf:
    bad = zf.testzip()
    xml = zf.read("word/document.xml").decode("utf-8")

summary = {
    "bytes": p.stat().st_size,
    "paragraphs": len(d.paragraphs),
    "tables": len(d.tables),
    "headings": len(headings),
    "images": len([r for r in d.part.rels.values() if "image" in r.reltype]),
    "explicit_page_breaks": xml.count('w:type="page"'),
    "hyperlinks": xml.count("<w:hyperlink"),
    "alt_text_present": "AR 四层主链" in xml,
    "zip_bad_member": bad,
    "product_plan_present": "未来产品落地计划" in all_text,
    "eric_present": "Eric" in all_text,
    "latest_commit_present": "5edc78e8" in all_text,
    "open_pr_count_present": "17 个" in all_text or "17 个 PR" in all_text,
}
print(summary)
print("FIRST_HEADINGS", headings[:12])
print("LAST_HEADINGS", headings[-8:])

assert bad is None
assert summary["tables"] == 34
assert summary["images"] >= 1
assert summary["alt_text_present"]
assert summary["product_plan_present"]
assert summary["eric_present"]
assert summary["latest_commit_present"]
assert summary["open_pr_count_present"]
